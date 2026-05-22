"""
acm_report.py — Generador de informe ACM en formato .docx

Cambios v2:
- Firma: recibe rango_min y rango_max en vez de precio_final único
- Sección 2 limpia: solo mediana, rango Q1-Q3, universo, días en mercado, señal
- Sin Top 3, sin radio, sin stock 0, sin ajuste por estado, sin precio sugerido
- Tabla de comparables: sin Dist., sin Piso → agrega Cochera, Antigüedad, Días pub, Expensas
- Nueva sección 3: "Tu propiedad en el mercado" (posicionamiento vs universo)
- Sección 4: Rango de publicación recomendado (los dos valores editables)
"""

from datetime import date
from io import BytesIO

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

COLOR_PRIMARY = RGBColor(0x00, 0x31, 0x53)
COLOR_ORANGE  = RGBColor(0xE8, 0x63, 0x0A)
COLOR_WHITE   = RGBColor(0xFF, 0xFF, 0xFF)
COLOR_GRAY    = RGBColor(0x6B, 0x72, 0x80)
COLOR_GREEN   = RGBColor(0x16, 0xA3, 0x4A)

_MESES = {
    "January": "enero", "February": "febrero", "March": "marzo",
    "April": "abril", "May": "mayo", "June": "junio",
    "July": "julio", "August": "agosto", "September": "septiembre",
    "October": "octubre", "November": "noviembre", "December": "diciembre",
}


# ─── Helpers de formato ────────────────────────────────────────────────────────

def fmt_usd(n) -> str:
    if n is None or n == 0:
        return "—"
    try:
        return f"USD {int(n):,}".replace(",", ".")
    except (TypeError, ValueError):
        return "—"


def fmt_ars(n) -> str:
    if n is None or n == 0:
        return "—"
    try:
        return f"$ {int(n):,}".replace(",", ".")
    except (TypeError, ValueError):
        return "—"


def fecha_es() -> str:
    hoy = date.today()
    mes_en = hoy.strftime("%B")
    return hoy.strftime(f"%d de {_MESES.get(mes_en, mes_en)} de %Y")


# ─── Helpers docx ─────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)


def _cell_text(cell, text: str, bold=False, color=None, size=10,
               align=None, italic=False):
    cell.text = ""
    para = cell.paragraphs[0]
    if align is not None:
        para.alignment = align
    run = para.add_run(str(text))
    run.bold        = bold
    run.italic      = italic
    run.font.size   = Pt(size)
    run.font.name   = "Calibri"
    if color:
        run.font.color.rgb = color


def _section_title(doc: Document, text: str, space_before: float = 14):
    p   = doc.add_paragraph()
    p.clear()
    run = p.add_run(text)
    run.bold              = True
    run.font.size         = Pt(13)
    run.font.name         = "Calibri"
    run.font.color.rgb    = COLOR_PRIMARY
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    return p


def _orange_line(doc: Document):
    p    = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "8")
    bot.set(qn("w:color"), "E8630A")
    pBdr.append(bot)
    pPr.append(pBdr)


def _add_hyperlink(paragraph, url: str, text: str):
    try:
        part  = paragraph.part
        r_id  = part.relate_to(
            url,
            "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
            is_external=True,
        )
        hl    = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), r_id)
        wr    = OxmlElement("w:r")
        rPr   = OxmlElement("w:rPr")
        col   = OxmlElement("w:color");  col.set(qn("w:val"), "E8630A");  rPr.append(col)
        u     = OxmlElement("w:u");      u.set(qn("w:val"), "single");    rPr.append(u)
        fnt   = OxmlElement("w:rFonts"); fnt.set(qn("w:ascii"), "Calibri"); rPr.append(fnt)
        sz    = OxmlElement("w:sz");     sz.set(qn("w:val"), "18");       rPr.append(sz)
        wr.append(rPr)
        t     = OxmlElement("w:t");      t.text = text
        wr.append(t)
        hl.append(wr)
        paragraph._p.append(hl)
    except Exception:
        paragraph.add_run(text)


def _set_col_widths(table, widths_cm: list):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


# ─── Función principal ────────────────────────────────────────────────────────

def generar_docx(
    datos_propiedad: dict,
    comparables: list,
    stats: dict,
    top3_indices: list,   # se mantiene en la firma por compatibilidad, ya no se usa
    precio_final: float,  # compatibilidad — se usa como referencia interna
    rango_min: float = 0,
    rango_max: float = 0,
) -> bytes:

    doc = Document()

    # A4, márgenes 2cm
    for section in doc.sections:
        section.page_height   = Cm(29.7)
        section.page_width    = Cm(21.0)
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.0)
        section.right_margin  = Cm(2.0)

        footer = section.footer
        fp     = footer.paragraphs[0]
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("Nahuel Lim · Asesor Inmobiliario · nahuelim.com.ar")
        r1.font.size = Pt(8); r1.font.name = "Calibri"; r1.font.color.rgb = COLOR_GRAY
        fp.add_run("\n")
        r2 = fp.add_run(
            "Valores de oferta publicada. No representan precios de cierre. "
            "La información es orientativa y no constituye asesoramiento legal ni financiero."
        )
        r2.font.size = Pt(7); r2.font.name = "Calibri"
        r2.font.color.rgb = COLOR_GRAY; r2.italic = True

    # ── HEADER ──────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    r = p.add_run("Análisis Comparativo de Mercado")
    r.bold = True; r.font.size = Pt(22); r.font.name = "Calibri"
    r.font.color.rgb = COLOR_PRIMARY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run("Nahuel Lim · Asesor Inmobiliario")
    r.font.size = Pt(11); r.font.name = "Calibri"; r.font.color.rgb = COLOR_GRAY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    r = p.add_run(f"Fecha: {fecha_es()}")
    r.font.size = Pt(10); r.font.name = "Calibri"; r.font.color.rgb = COLOR_GRAY
    p.paragraph_format.space_after = Pt(6)

    _orange_line(doc)

    # ── SECCIÓN 1: Propiedad Analizada ──────────────────────────────────────────
    _section_title(doc, "1. Propiedad Analizada")

    campos = [
        ("Dirección",   datos_propiedad.get("direccion") or "—"),
        ("Barrio",      datos_propiedad.get("barrio")    or "—"),
        ("Tipo",        datos_propiedad.get("tipo")      or "—"),
        ("Ambientes",   str(datos_propiedad.get("ambientes") or "—")),
        ("M² cubiertos", f'{datos_propiedad.get("sup_cubierta") or "—"} m²'),
    ]
    if datos_propiedad.get("sup_semicubierta") and \
       float(datos_propiedad.get("sup_semicubierta", 0) or 0) > 0:
        campos.append(("M² semicubiertos",
                        f'{datos_propiedad["sup_semicubierta"]} m²'))
    if datos_propiedad.get("antiguedad"):
        campos.append(("Antigüedad", f'{datos_propiedad["antiguedad"]} años'))
    if datos_propiedad.get("piso"):
        campos.append(("Piso", str(datos_propiedad["piso"])))
    if datos_propiedad.get("disposicion"):
        campos.append(("Disposición", datos_propiedad["disposicion"]))
    campos.append(("Cochera",      "Sí" if datos_propiedad.get("cochera") else "No"))
    if datos_propiedad.get("expensas"):
        campos.append(("Expensas est.", fmt_ars(datos_propiedad["expensas"])))
    if datos_propiedad.get("estado"):
        campos.append(("Estado del inmueble", datos_propiedad["estado"]))

    t1 = doc.add_table(rows=1, cols=2)
    t1.style = "Table Grid"
    _set_cell_bg(t1.rows[0].cells[0], "003153")
    _set_cell_bg(t1.rows[0].cells[1], "003153")
    _cell_text(t1.rows[0].cells[0], "Campo", bold=True, color=COLOR_WHITE, size=10)
    _cell_text(t1.rows[0].cells[1], "Valor", bold=True, color=COLOR_WHITE, size=10)
    for i, (campo, valor) in enumerate(campos):
        row = t1.add_row()
        bg  = "EEF3F8" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(row.cells[0], bg); _set_cell_bg(row.cells[1], bg)
        _cell_text(row.cells[0], campo, bold=True, size=10)
        _cell_text(row.cells[1], valor, size=10)

    doc.add_paragraph()

    # ── SECCIÓN 2: Análisis de Mercado ──────────────────────────────────────────
    _section_title(doc, "2. Análisis de Mercado")

    mediana_m2  = stats.get("mediana_m2", 0)
    p25         = stats.get("p25", 0)
    p75         = stats.get("p75", 0)
    n_validos   = stats.get("n_validos", 0)
    n_total     = stats.get("n_total", 0)
    dias_med    = stats.get("dias_mediana")
    pct_30      = stats.get("dias_pct_30", 0)
    pct_90      = stats.get("dias_pct_90", 0)
    pct_mas90   = stats.get("dias_pct_mas90", 0)
    señal       = stats.get("señal_mercado", "")

    rango_mkt = (
        f"USD {p25:,.0f} – USD {p75:,.0f}".replace(",", ".")
        if p25 and p75 else "—"
    )

    stats_rows = [
        ("Propiedades analizadas",
         f"{n_validos} propiedades" + (f" (de {n_total} relevadas)" if n_total > n_validos else "")),
        ("Mediana precio/m²",          fmt_usd(mediana_m2)),
        ("Rango de mercado (Q1–Q3)",   rango_mkt),
    ]
    if dias_med is not None:
        stats_rows += [
            ("Mediana días en mercado", f"{dias_med} días"),
            ("Stock con ≤ 30 días",    f"{pct_30}%  — rotación activa"),
            ("Stock con 31–90 días",   f"{pct_90}%"),
            ("Stock con > 90 días",    f"{pct_mas90}%  — sin movimiento"),
        ]

    t2 = doc.add_table(rows=1, cols=2)
    t2.style = "Table Grid"
    _set_cell_bg(t2.rows[0].cells[0], "003153")
    _set_cell_bg(t2.rows[0].cells[1], "003153")
    _cell_text(t2.rows[0].cells[0], "Indicador", bold=True, color=COLOR_WHITE, size=10)
    _cell_text(t2.rows[0].cells[1], "Valor",     bold=True, color=COLOR_WHITE, size=10)
    for i, (label, valor) in enumerate(stats_rows):
        row = t2.add_row()
        bg  = "EEF3F8" if i % 2 == 0 else "FFFFFF"
        _set_cell_bg(row.cells[0], bg); _set_cell_bg(row.cells[1], bg)
        _cell_text(row.cells[0], label, bold=True, size=10)
        _cell_text(row.cells[1], valor, size=10)

    # Señal de mercado como párrafo destacado
    if señal:
        doc.add_paragraph()
        p_señal = doc.add_paragraph()
        p_señal.paragraph_format.left_indent  = Cm(0.5)
        p_señal.paragraph_format.right_indent = Cm(0.5)
        p_señal.paragraph_format.space_before = Pt(4)
        p_señal.paragraph_format.space_after  = Pt(4)
        r_señal = p_señal.add_run(f"📊  {señal}")
        r_señal.font.size      = Pt(9.5)
        r_señal.font.name      = "Calibri"
        r_señal.font.color.rgb = COLOR_PRIMARY
        r_señal.italic         = True

    doc.add_paragraph()

    # ── SECCIÓN 3: Tu propiedad en el mercado ────────────────────────────────────
    _section_title(doc, "3. Tu propiedad en el mercado")

    sup = float(datos_propiedad.get("sup_cubierta") or 0)

    # Calcular cuántas propiedades están por debajo, en rango y por encima del rango
    r_lo = rango_min if rango_min else (p25 * sup if p25 and sup else 0)
    r_hi = rango_max if rango_max else (p75 * sup if p75 and sup else 0)

    n_por_debajo = sum(1 for c in comparables if c.get("precio") and c["precio"] < r_lo)
    n_en_rango   = sum(1 for c in comparables if c.get("precio") and r_lo <= c["precio"] <= r_hi)
    n_por_encima = sum(1 for c in comparables if c.get("precio") and c["precio"] > r_hi)
    n_comp       = len(comparables)

    # Posicionamiento por precio/m² vs mediana
    rango_m2_lo = r_lo / sup if sup > 0 and r_lo > 0 else 0
    rango_m2_hi = r_hi / sup if sup > 0 and r_hi > 0 else 0

    pos_rows = []
    if n_comp > 0:
        pos_rows.append((
            "Propiedades que compiten directamente",
            f"{n_comp} publicaciones activas en el mercado"
        ))
    if r_lo and r_hi:
        pos_rows.append((
            "Por debajo del rango recomendado",
            f"{n_por_debajo} propiedades ({int(n_por_debajo/n_comp*100) if n_comp else 0}%)"
        ))
        pos_rows.append((
            "Dentro del rango recomendado",
            f"{n_en_rango} propiedades ({int(n_en_rango/n_comp*100) if n_comp else 0}%)"
        ))
        pos_rows.append((
            "Por encima del rango recomendado",
            f"{n_por_encima} propiedades ({int(n_por_encima/n_comp*100) if n_comp else 0}%)"
        ))
    if mediana_m2 and rango_m2_lo and rango_m2_hi:
        if rango_m2_lo <= mediana_m2 <= rango_m2_hi:
            posicion = "dentro de la mediana de mercado"
        elif rango_m2_hi < mediana_m2:
            posicion = f"por debajo de la mediana (USD {int(mediana_m2):,}/m²)".replace(",",".")
        else:
            posicion = f"por encima de la mediana (USD {int(mediana_m2):,}/m²)".replace(",",".")
        pos_rows.append(("Posicionamiento vs. mediana", posicion))

    if pos_rows:
        t3 = doc.add_table(rows=1, cols=2)
        t3.style = "Table Grid"
        _set_cell_bg(t3.rows[0].cells[0], "003153")
        _set_cell_bg(t3.rows[0].cells[1], "003153")
        _cell_text(t3.rows[0].cells[0], "Indicador", bold=True, color=COLOR_WHITE, size=10)
        _cell_text(t3.rows[0].cells[1], "Valor",     bold=True, color=COLOR_WHITE, size=10)
        for i, (label, valor) in enumerate(pos_rows):
            row = t3.add_row()
            bg  = "EEF3F8" if i % 2 == 0 else "FFFFFF"
            _set_cell_bg(row.cells[0], bg); _set_cell_bg(row.cells[1], bg)
            _cell_text(row.cells[0], label, bold=True, size=10)
            _cell_text(row.cells[1], valor, size=10)

    doc.add_paragraph()

    # ── SECCIÓN 4: Rango de Publicación Recomendado ──────────────────────────────
    _section_title(doc, "4. Rango de Publicación Recomendado")

    rango_m2_txt = ""
    if rango_m2_lo and rango_m2_hi:
        rango_m2_txt = (
            f"USD {int(rango_m2_lo):,} – USD {int(rango_m2_hi):,} /m²"
        ).replace(",", ".")

    rango_rows = [
        ("Rango recomendado",
         f"{fmt_usd(r_lo)} – {fmt_usd(r_hi)}" if r_lo and r_hi else "—"),
        ("Precio/m² del rango",
         rango_m2_txt or "—"),
        ("Precio de referencia interno",
         fmt_usd(precio_final) if precio_final else "—"),
    ]

    t4 = doc.add_table(rows=1, cols=2)
    t4.style = "Table Grid"
    _set_cell_bg(t4.rows[0].cells[0], "003153")
    _set_cell_bg(t4.rows[0].cells[1], "003153")
    _cell_text(t4.rows[0].cells[0], "Campo",  bold=True, color=COLOR_WHITE, size=10)
    _cell_text(t4.rows[0].cells[1], "Valor",  bold=True, color=COLOR_WHITE, size=10)
    for i, (label, valor) in enumerate(rango_rows):
        row = t4.add_row()
        is_rango = i == 0
        bg = "E8F4ED" if is_rango else ("EEF3F8" if i % 2 == 0 else "FFFFFF")
        _set_cell_bg(row.cells[0], bg); _set_cell_bg(row.cells[1], bg)
        _cell_text(row.cells[0], label, bold=True,   size=10,
                   color=COLOR_GREEN if is_rango else None)
        _cell_text(row.cells[1], valor, bold=is_rango, size=11 if is_rango else 10,
                   color=COLOR_GREEN if is_rango else None)

    doc.add_paragraph()

    # ── SECCIÓN 5: Testigos de Mercado ───────────────────────────────────────────
    _section_title(doc, "5. Testigos de Mercado")

    p_intro = doc.add_paragraph()
    r_intro = p_intro.add_run(
        f"Los {len(comparables)} comparables seleccionados a continuación son las propiedades "
        "más similares activas en el mercado al momento de este análisis."
    )
    r_intro.font.size = Pt(9.5); r_intro.font.name = "Calibri"
    r_intro.font.color.rgb = COLOR_GRAY
    p_intro.paragraph_format.space_after = Pt(8)

    if comparables:
        # Columnas: Dirección | Amb | m² | Coch | Antigüedad | Días pub | USD/m² | Precio | Expensas | Link
        col_headers = [
            "Dirección", "Amb", "m²", "Coch",
            "Antigüedad", "Días pub", "USD/m²", "Precio", "Expensas", "Ver"
        ]
        widths = [4.5, 0.8, 0.8, 0.8, 1.2, 1.0, 1.4, 1.8, 1.8, 0.9]

        t5 = doc.add_table(rows=1, cols=len(col_headers))
        t5.style = "Table Grid"
        hdr5 = t5.rows[0].cells
        for i, h in enumerate(col_headers):
            _set_cell_bg(hdr5[i], "003153")
            alg = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
            _cell_text(hdr5[i], h, bold=True, color=COLOR_WHITE, size=8, align=alg)

        for j, comp in enumerate(comparables):
            row = t5.add_row()
            bg  = "F4F4F4" if j % 2 == 0 else "FFFFFF"
            for cell in row.cells:
                _set_cell_bg(cell, bg)

            dir_txt  = (comp.get("direccion") or comp.get("barrio") or "—")[:42]
            amb_txt  = str(comp.get("ambientes") or "—")
            sup_txt  = str(int(comp["sup_cubierta"])) if comp.get("sup_cubierta") else "—"
            coch_txt = "Sí" if comp.get("cochera") else "No"
            antig    = f"{int(comp['antiguedad'])} años" if comp.get("antiguedad") else "—"
            dias     = comp.get("dias_publicado")
            if dias is None:
                dias_txt = "—"
            elif dias <= 30:
                dias_txt = f"{dias}d ✓"
            elif dias <= 90:
                dias_txt = f"{dias}d"
            else:
                dias_txt = f"{dias}d !"
            m2_txt   = fmt_usd(int(comp["precio_por_m2"])) if comp.get("precio_por_m2") else "—"
            prec_txt = fmt_usd(int(comp["precio"]))         if comp.get("precio")       else "—"
            exp_txt  = fmt_ars(comp["expensas"])             if comp.get("expensas")     else "—"

            C = WD_ALIGN_PARAGRAPH.CENTER
            _cell_text(row.cells[0], dir_txt,  size=8)
            _cell_text(row.cells[1], amb_txt,  size=8, align=C)
            _cell_text(row.cells[2], sup_txt,  size=8, align=C)
            _cell_text(row.cells[3], coch_txt, size=8, align=C)
            _cell_text(row.cells[4], antig,    size=8, align=C)
            _cell_text(row.cells[5], dias_txt, size=8, align=C)
            _cell_text(row.cells[6], m2_txt,   size=8, align=C)
            _cell_text(row.cells[7], prec_txt, size=8, align=C)
            _cell_text(row.cells[8], exp_txt,  size=8, align=C)

            row.cells[9].text = ""
            p_link = row.cells[9].paragraphs[0]
            p_link.alignment = C
            if comp.get("url"):
                _add_hyperlink(p_link, comp["url"], "Ver →")

        _set_col_widths(t5, widths)

        # Leyenda días
        p_ley = doc.add_paragraph()
        r_ley = p_ley.add_run(
            "Días pub: ✓ ≤ 30 días (rotación activa)  ·  ! > 90 días (sin movimiento)"
        )
        r_ley.font.size = Pt(7.5); r_ley.font.color.rgb = COLOR_GRAY; r_ley.italic = True

    doc.add_paragraph()

    # ── Guardar ──────────────────────────────────────────────────────────────────
    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
